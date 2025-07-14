from pydantic import BaseModel
from io import BytesIO
import pdfplumber
from fpdf import FPDF
from docx import Document

class FileMCQRequest(BaseModel):
    num_questions: int

# --------- Utils ----------
def extract_text(file_path: str, file_bytes: BytesIO) -> str:
    ext = file_path.rsplit('.', 1)[-1].lower()
    if ext == "pdf":
        with pdfplumber.open(file_bytes) as pdf:
            return ''.join([p.extract_text() for p in pdf.pages if p.extract_text()])
    elif ext == "docx":
        doc = Document(file_bytes)
        return ' '.join([para.text for para in doc.paragraphs])
    elif ext == "txt":
        return file_bytes.read().decode('utf-8')
    else:
        raise ValueError("Unsupported file type")


def generate_mcq_docs(mcqs: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for mcq in mcqs.split("## MCQ"):
        pdf.multi_cell(0, 10, mcq.strip())
    pdf_bytes = pdf.output(dest='S').encode('latin-1')

    docx_file = BytesIO()
    doc = Document()
    for mcq in mcqs.split("## MCQ"):
        doc.add_paragraph(mcq.strip())
    doc.save(docx_file)
    docx_file.seek(0)

    return pdf_bytes, docx_file