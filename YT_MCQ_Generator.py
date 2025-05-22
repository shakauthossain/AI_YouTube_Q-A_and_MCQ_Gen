import streamlit as st
import textwrap
from datetime import timedelta
from io import BytesIO
import os
import pdfplumber
import docx
from fpdf import FPDF
from docx import Document

from langchain.document_loaders import YoutubeLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import LLMChain
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.prompts.chat import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate

from dotenv import load_dotenv

# Page setup
st.set_page_config(page_title="YouTube & File MCQ Generator", layout="centered")
st.title("YouTube + File-Based Q&A and MCQ Generator")

load_dotenv()
api_key = os.getenv("GROQ_API")

# Initialize LLM
chat = ChatGroq(
    api_key=api_key,  # Replace with your actual key
    model="llama-3.3-70b-versatile",
    temperature=0.5
)

# Prompt template for MCQs
mcq_prompt = PromptTemplate(
    input_variables=["context", "num_questions"],
    template="""
You are an AI assistant helping the user generate multiple-choice questions (MCQs) from the text below:

Text:
{context}

Generate {num_questions} MCQs. Each should include:
- A clear question
- Four answer options labeled A, B, C, and D
- The correct answer clearly indicated at the end

Format:
## MCQ
Question: [question]
A) [option A]
B) [option B]
C) [option C]
D) [option D]
Correct Answer: [correct option]
"""
)

# Function to extract text from file
def extract_text(file_path, file_bytes):
    ext = file_path.rsplit('.', 1)[-1].lower()
    if ext == "pdf":
        with pdfplumber.open(file_bytes) as pdf:
            return ''.join([p.extract_text() for p in pdf.pages if p.extract_text()])
    elif ext == "docx":
        doc = docx.Document(file_bytes)
        return ' '.join([para.text for para in doc.paragraphs])
    elif ext == "txt":
        return file_bytes.read().decode('utf-8')
    else:
        raise ValueError("Unsupported file type")

# Create tabs
tab1, tab2 = st.tabs(["YouTube Q&A + MCQ Generator", "File-Based MCQ Generator"])

# ---------------- Tab 1: YouTube Q&A + MCQs ----------------
with tab1:
    st.subheader("1. Load YouTube Transcript and Ask Questions")
    if "db" not in st.session_state:
        st.session_state.db = None
    if "docs" not in st.session_state:
        st.session_state.docs = None

    col1, col2 = st.columns([6, 1])
    with col1:
        video_url = st.text_input("YouTube Video URL", label_visibility="collapsed", placeholder="https://www.youtube.com/watch?v=...")
    with col2:
        load_clicked = st.button("Load", use_container_width=True)

    if load_clicked and video_url:
        with st.spinner("Loading transcript..."):
            try:
                loader = YoutubeLoader.from_youtube_url(video_url)
                transcript = loader.load()
                splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)
                docs = splitter.split_documents(transcript)
                db = FAISS.from_documents(docs, HuggingFaceEmbeddings())
                st.session_state.db = db
                st.session_state.docs = docs
                st.success("Transcript loaded successfully.")
            except Exception as e:
                st.error(f"Failed to load transcript: {e}")

    st.divider()
    st.subheader("2. Ask a Question from Transcript")

    col3, col4 = st.columns([6, 1])
    with col3:
        query = st.text_input("Enter your question", label_visibility="collapsed", placeholder="e.g., What is the video about?")
    with col4:
        answer_clicked = st.button("Answer", use_container_width=True)

    if answer_clicked:
        if not query:
            st.warning("Please enter a question.")
        elif not st.session_state.db:
            st.warning("Please load a transcript first.")
        else:
            with st.spinner("Finding answer..."):
                results = st.session_state.db.similarity_search(query, k=4)
                docs_page_content = " ".join([doc.page_content for doc in results])

                system_template = """You are a helpful assistant answering questions about a YouTube video transcript.
Use only the factual information from the transcript to answer the question.

Transcript: {docs}
If unsure, say "I don't know."
"""
                human_template = "Answer the following question: {question}"

                prompt = ChatPromptTemplate.from_messages([
                    SystemMessagePromptTemplate.from_template(system_template),
                    HumanMessagePromptTemplate.from_template(human_template)
                ])
                chain = LLMChain(llm=chat, prompt=prompt)
                answer = chain.run(question=query, docs=docs_page_content)
                st.success("Answer:")
                st.write(textwrap.fill(answer.strip(), width=80))

                with st.expander("Transcript Snippets Used"):
                    for i, doc in enumerate(results):
                        meta = doc.metadata
                        if 'start' in meta:
                            t = str(timedelta(seconds=int(float(meta['start']))))
                            link = f"{video_url}&t={int(float(meta['start']))}s"
                            st.markdown(f"**Snippet {i+1} at [`{t}`]({link}):**")
                        else:
                            st.markdown(f"**Snippet {i+1}:**")
                        st.info(textwrap.fill(doc.page_content, width=80))

    st.divider()
    st.subheader("3. Generate MCQs from YouTube Transcript")
    col5, col6 = st.columns([6, 1])
    with col5:
        num_mcqs_youtube = st.slider("Number of MCQs", 1, 10, 5, key="slider_youtube", label_visibility="collapsed")
    with col6:
        generate_mcqs = st.button("Generate MCQs")

    if generate_mcqs:
        if not st.session_state.docs:
            st.warning("Please load the transcript first.")
        else:
            with st.spinner("Generating MCQs..."):
                full_text = " ".join([doc.page_content for doc in st.session_state.docs])
                chain = LLMChain(llm=chat, prompt=mcq_prompt)
                mcqs = chain.run(context=full_text, num_questions=num_mcqs_youtube).strip()

                st.success("MCQs Generated:")
                for mcq in mcqs.split("## MCQ"):
                    if mcq.strip():
                        st.markdown(f"```{mcq.strip()}```")

                # PDF + DOCX
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for mcq in mcqs.split("## MCQ"):
                    pdf.multi_cell(0, 10, mcq.strip())
                pdf_bytes = pdf.output(dest='S').encode('latin-1')

                docx_file = BytesIO()
                doc = Document()
                for mcq in mcqs.split("## MCQ"):
                    doc.add_paragraph(mcq.strip())
                doc.save(docx_file)
                docx_file.seek(0)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📄 Download PDF", pdf_bytes, "youtube_mcqs.pdf", "application/pdf")
                with col2:
                    st.download_button("📝 Download DOCX", docx_file, "youtube_mcqs.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------------- Tab 2: File Upload MCQs ----------------
with tab2:
    st.subheader("Upload File and Generate MCQs")
    uploaded_file = st.file_uploader("Upload TXT, PDF, or DOCX file", type=["txt", "pdf", "docx"])
    num_mcqs_file = st.slider("Number of MCQs", 1, 10, 5, key="slider_file")
    if st.button("Generate File MCQs"):
        if not uploaded_file:
            st.warning("Please upload a file.")
        else:
            try:
                with st.spinner("Extracting text..."):
                    content = extract_text(uploaded_file.name, uploaded_file)

                with st.spinner("Generating MCQs..."):
                    chain = LLMChain(llm=chat, prompt=mcq_prompt)
                    mcqs = chain.run(context=content, num_questions=num_mcqs_file).strip()

                    st.success("MCQs Generated:")
                    for mcq in mcqs.split("## MCQ"):
                        if mcq.strip():
                            st.markdown(f"```{mcq.strip()}```")

                    # PDF + DOCX
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    for mcq in mcqs.split("## MCQ"):
                        pdf.multi_cell(0, 10, mcq.strip())
                    pdf_bytes = pdf.output(dest='S').encode('latin-1')

                    docx_file = BytesIO()
                    doc = Document()
                    for mcq in mcqs.split("## MCQ"):
                        doc.add_paragraph(mcq.strip())
                    doc.save(docx_file)
                    docx_file.seek(0)

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button("📄 Download PDF", pdf_bytes, "file_mcqs.pdf", "application/pdf")
                    with col2:
                        st.download_button("📝 Download DOCX", docx_file, "file_mcqs.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"Failed to process file: {e}")
